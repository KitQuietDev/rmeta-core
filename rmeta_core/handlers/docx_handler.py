# handlers/docx_handler.py

"""
DOCX Metadata Scrubber for rMeta

Copies paragraph content to a fresh document to remove embedded metadata.
Format: .docx
Non-destructive to content
"""

import logging
import os
from pathlib import Path
import asyncio
import docx
from rmeta_core.utils.pii_scanner import scan_text_for_pii

logger = logging.getLogger(__name__)
__all__ = ["scrub", "get_additional_messages"]

# Supported file types for this handler
SUPPORTED_EXTENSIONS = {"docx"}

# Indicates this handler supports PII detection via scan_text_for_pii
PII_DETECT = True

async def scrub(file_path: str) -> None:
    path = Path(file_path)
    ext = path.suffix.lower().lstrip(".")

    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
    if not os.access(file_path, os.R_OK | os.W_OK):
        raise PermissionError(f"Cannot access DOCX file: {file_path}")
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    # Make the nested function sync since python-docx is sync
    def scrub_docx():
        doc = docx.Document(file_path)
        new_doc = docx.Document()
        for para in doc.paragraphs:
            new_doc.add_paragraph(para.text)
        temp_path = path.with_suffix(".tmp.docx")
        new_doc.save(temp_path) # type: ignore
        os.replace(temp_path, path)
        logger.info(f"DOCX scrubbed: {file_path}")

    # Run the sync function in a thread pool
    await asyncio.to_thread(scrub_docx)

    if not path.exists() or path.stat().st_size == 0:
        raise RuntimeError(f"Scrubbed DOCX file missing or empty: {file_path}")

async def get_additional_messages(file_path: str) -> list[str]:
    messages = [f"Metadata stripped from DOCX: {Path(file_path).name}"]

    def extract_and_scan():
        doc = docx.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return scan_text_for_pii(text)

    try:
        pii_found = await asyncio.to_thread(extract_and_scan)
        for pii_type in pii_found:
            messages.append(f"PII detected: {pii_type.title()} found in file")
    except Exception as e:
        logger.warning(f"Could not scan DOCX for PII: {e}")

    return messages