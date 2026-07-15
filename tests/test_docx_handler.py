import asyncio

import docx
import pytest

from rmeta_core.handlers.docx_handler import scrub, get_additional_messages


def test_scrub_resets_properties_preserves_text(dirty_docx):
    before = docx.Document(str(dirty_docx))
    assert before.core_properties.author == "MetaMaker"
    original_text = before.paragraphs[0].text

    asyncio.run(scrub(str(dirty_docx)))

    after = docx.Document(str(dirty_docx))
    assert after.core_properties.author != "MetaMaker"
    assert after.paragraphs[0].text == original_text


def test_get_additional_messages_flags_pii(dirty_docx):
    asyncio.run(scrub(str(dirty_docx)))
    messages = asyncio.run(get_additional_messages(str(dirty_docx)))
    assert any("PII detected" in m for m in messages)


def test_scrub_missing_file_raises():
    with pytest.raises(FileNotFoundError):
        asyncio.run(scrub("does-not-exist.docx"))
